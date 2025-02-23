from flask import Flask, render_template, request, jsonify, send_file, url_for, send_from_directory
import openai, whisper, os, re
from dotenv import load_dotenv
from docx import Document 
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from weasyprint import HTML
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__, static_folder='static')
whisper_model = whisper.load_model("small")
os.makedirs('uploads', exist_ok=True)


def parse_cv_text(text):
    """Parse CV sections from raw text."""
    return {k: (v.group(1).strip() if v else "") for k, v in {
        "Name": re.search(r"Name[:\-]\s*(.*)", text, re.IGNORECASE),
        "Title": re.search(r"Title[:\-]\s*(.*)", text, re.IGNORECASE),
    }.items()}

def generate_cv_pdf(data):
    """Generate CV in PDF format."""
    output_path = os.path.join('uploads', 'cv_result.pdf')
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height - 50, data.get("Name", ""))
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 75, data.get("Title", ""))

    y = height - 110
    for section, content in data.items():
        if section not in ["Name", "Title"] and content:
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, section)
            c.setFont("Helvetica", 11)
            text_obj = c.beginText(60, y - 15)
            for line in content.split('\n'):
                text_obj.textLine(line)
            c.drawText(text_obj)
            y -= (15 * (len(content.split('\n')) + 2))
    c.save()
    return output_path

def generate_docx(data, output_filename):
    """Generate a DOCX CV."""
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)
    for section, content in data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content if isinstance(content, str) else "\n".join(content))
    doc.save(output_filename)
    return output_filename

def docx_to_html(docx_path):
    """Convert DOCX to HTML using LibreOffice."""
    output_html_path = docx_path.replace('.docx', '.html')
    os.system(f"/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to html {docx_path}  --outdir uploads")
    return output_html_path

def parse_cv_text(text):
    """Parse CV sections from raw text."""
    return {k: (v.group(1).strip() if v else "") for k, v in {
        "Name": re.search(r"Name[:\-]\s*(.*)", text, re.IGNORECASE),
        "Title": re.search(r"Title[:\-]\s*(.*)", text, re.IGNORECASE),
    }.items()}

def generate_cv_pdf(data):
    """Generate CV in PDF format."""
    output_path = os.path.join('uploads', 'cv_result.pdf')
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4

    c.setFont("Helvetica-Bold", 20)
    c.drawString(50, height - 50, data.get("Name", ""))
    c.setFont("Helvetica", 14)
    c.drawString(50, height - 75, data.get("Title", ""))

    y = height - 110
    for section, content in data.items():
        if section not in ["Name", "Title"] and content:
            c.setFont("Helvetica-Bold", 12)
            c.drawString(50, y, section)
            c.setFont("Helvetica", 11)
            text_obj = c.beginText(60, y - 15)
            for line in content.split('\n'):
                text_obj.textLine(line)
            c.drawText(text_obj)
            y -= (15 * (len(content.split('\n')) + 2))
    c.save()
    return output_path

def generate_docx(data, output_filename):
    """Generate a DOCX CV."""
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)
    for section, content in data.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content if isinstance(content, str) else "\n".join(content))
    doc.save(output_filename)
    return output_filename

def generate_enhanced_docx(data, output_filename):
    doc = Document()

    # Get the name from the data dictionary and use it as the title
    name = data.get("Name", "Curriculum Vitae")  # Default to "Curriculum Vitae" if Name is not found

    # Add a title heading with large font size
    title = doc.add_heading(level=0)
    run = title.add_run(name)  # Use the Name field as the title
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centered title
    doc.add_paragraph()  # Add a blank line

    # Iterate through the sections and format them nicely
    for section, content in data.items():
        if section == "Name":
            continue  # Skip the Name section as it's already used as the title

        # Add section header (with a different style)
        heading = doc.add_heading(level=1)
        heading_run = heading.add_run(section)
        heading_run.font.size = Pt(16)
        heading_run.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for headings

        doc.add_paragraph()  # Add space after heading
        
        # Handle content based on its type
        if isinstance(content, list):
            # If the section content is a list, format it as bullet points
            for item in content:
                doc.add_paragraph(f"- {item}", style='List Bullet')
        
        elif isinstance(content, dict):
            # If content is a dictionary, create a table dynamically
            table = doc.add_table(rows=1, cols=len(content))
            table.style = 'Table Grid'
            row = table.rows[0].cells
            for i, key in enumerate(content.keys()):
                row[i].text = key  # Set the headers (keys)
            row = table.add_row().cells
            for i, value in enumerate(content.values()):
                row[i].text = value  # Set the values
            doc.add_paragraph()  # Add space after table

        elif isinstance(content, str):
            # If content is a string (text), add it as a normal paragraph
            paragraph = doc.add_paragraph(content)
            paragraph.style.font.size = Pt(12)
        
        doc.add_paragraph()  # Add space between sections

    # Save the document
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")
    return output_filename

def docx_to_html(docx_path):
    """Convert DOCX to HTML using LibreOffice."""
    output_html_path = docx_path.replace('.docx', '.html')
    os.system(f"/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to html {docx_path}  --outdir uploads")
    return output_html_path

@app.route('/')
def home(): return render_template('index.html')

@app.route('/demo')
def demo(): return render_template('demo.html')

@app.route('/about')
def about(): return render_template('about.html')

@app.route('/features')
def features(): return render_template('features.html')

@app.route('/pricing')
def pricing(): return render_template('pricing.html')

@app.route('/contact')
def contact(): return render_template('contact.html')

# @app.route('/download')
# def download_cv():
#     return send_file('uploads/cv_result.pdf', as_attachment=True, download_name="Your_CV.pdf")

CV_QUESTIONS = {
    "Name": "What is your full name?",
    "Title": "What is your job title?",
    "Summary": "Provide a brief professional summary.",
    "Skills": "List your key skills.",
    "Experience": "Describe your work experience.",
    "Education": "Where did you study and what degree did you earn?",
    "Certifications": "Do you have any certifications or awards?"
}

# @app.route('/next_question', methods=['GET'])
# def next_question():
#     """🔎 Get the next CV question."""
#     index = int(request.args.get('index', 0))
#     keys = list(CV_QUESTIONS.keys())
#     return jsonify({"key": keys[index], "question": CV_QUESTIONS[keys[index]]}) if index < len(keys) else jsonify({"message": "All questions completed."})

@app.route('/transcribe', methods=['POST'])
def transcribe():
    """🎙️ Handle audio transcription with Whisper model."""
    if 'audio' not in request.files:
        return jsonify({'text': 'No file uploaded'}), 400

    audio_path = os.path.join('uploads', 'temp.wav')
    request.files['audio'].save(audio_path)

    try:
        result = whisper_model.transcribe(audio_path)
        transcript_text = result['text']
        structured_data = {request.form.get('question_key'): transcript_text}
        return jsonify({'text': transcript_text, 'structured_data': structured_data})
    except Exception as e:
        return jsonify({'text': f"Error: {str(e)}"}), 500
    finally:
        os.remove(audio_path)



@app.route('/improve', methods=['POST'])
def improve():
    """💎 Improve CV content with OpenAI."""
    text = request.json.get('text', '')
    if not text:
        return jsonify({'improved_text': 'No text provided.'}), 400

    prompt = f"Improve and structure the following CV content with professional language and clear sections:\n\n{text}"
    try:
        response = openai.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are an expert in creating professional CVs."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=800
        )
        improved_text = response.choices[0].message.content.strip()
        parsed_data = parse_cv_text(improved_text)
        generate_cv_pdf(parsed_data)
        return jsonify({'improved_text': improved_text, 'pdf_url': url_for('download_cv')})
    except Exception as e:
        return jsonify({'improved_text': f"Error: {str(e)}"}), 500

# @app.route('/download')
# def download_cv():
#     return send_file('uploads/cv_result.pdf', as_attachment=True, download_name="Your_CV.pdf")

CV_QUESTIONS = {
    "Name": "What is your full name?",
    "Title": "What is your job title?",
    "Summary": "Provide a brief professional summary.",
    "Skills": "List your key skills.",
    "Experience": "Describe your work experience.",
    "Education": "Where did you study and what degree did you earn?",
    "Certifications": "Do you have any certifications or awards?"
}

@app.route('/next_question', methods=['GET'])
def next_question():
    """Get the next question to be displayed on the frontend."""
    question_keys = list(CV_QUESTIONS.keys())
    current_index = int(request.args.get('index', 0))
    
    if current_index < len(question_keys):
        return jsonify({"key":question_keys[current_index],"question": CV_QUESTIONS[question_keys[current_index]], "index": current_index})
    else:
        return jsonify({"message": "All questions completed."})

def parse_transcript(question_key,text):
    """Extracts CV sections from a raw transcript using regex."""
    sections = {key: "" for key in CV_QUESTIONS.keys()}  # Initialize empty structure
    sections[question_key] = text;
    print(sections)
    return sections

# Function to generate a DOCX file based on structured CV data
# todo
def generate_docx(data, output_filename):
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)

    for section, content in data.items():
        doc.add_heading(section, level=1)
        if isinstance(content, list):
            for item in content:
                doc.add_paragraph(f"- {item}")
        else:
            doc.add_paragraph(content)

    doc.save(output_filename)
    print(f"Document saved as {output_filename}")
    return output_filename

# @app.route('/generate_docx', methods=['POST'])
# def generate_docx_file():
#     """Generate a DOCX file from structured CV data."""
#     data = request.json.get('structured_data', {})

#     if not data:
#         return jsonify({"error": "No structured data provided."}), 400

#     output_filename = "./uploads/output_cv.docx"
#     generate_docx(data, output_filename)

#     return jsonify({"message": "CV generated successfully", "file": output_filename, "html_file": docx_to_html(output_filename)})


# def docx_to_html(docx_path, output_html_path):
#     """Convert DOCX to HTML and save the result to a file."""
#     if not os.path.exists(docx_path):
#         raise FileNotFoundError(f"Input DOCX file not found: {docx_path}")

#     # Convert DOCX to HTML using LibreOffice headless mode
#     os.system(f"/Applications/LibreOffice.app/Contents/MacOS/soffice --headless --convert-to html {docx_path} --outdir ./uploads")

#     # # Check if the HTML file was created
#     # if not os.path.exists(output_html_path):
#     #     raise FileNotFoundError(f"Failed to generate HTML file: {output_html_path}")

#     return output_html_path



def html_to_pdf(html_path, pdf_path):
    # Read the HTML file
    html = HTML(filename=html_path)

    # Render HTML to PDF and save it
    html.write_pdf(pdf_path)

# Route to generate PDF from DOCX
@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    try:
        data = request.get_json()
        html_content = data.get('html')

        if not html_content:
            return jsonify({"error": "No HTML content provided"}), 400

        # Save the HTML content to a temporary file
        temp_html_path = "temp_document.html"
        with open(temp_html_path, "w", encoding="utfç-8") as file:
            file.write(html_content)

        # Convert the HTML file to a PDF
        pdf_path = "output.pdf"
        HTML(filename=temp_html_path).write_pdf(pdf_path)

        # Return the generated PDF file as a response
        return send_file(pdf_path, as_attachment=True)

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    

@app.route('/editor')
def editor():
    return render_template('editor.html') 

@app.route('/generate_html')
def generate_cv():
    docx_path = 'output_cv.docx'
    
    if os.path.exists(docx_path):
        html_file = docx_to_html(docx_path)

        # Read HTML and remove unwanted tags
        with open(html_file, 'r', encoding="utf-8") as file:
            html_content = file.read()

        # Remove unwanted <html>, <body>, etc.
        html_content = html_content.replace("<html>", "").replace("</html>", "")
        html_content = html_content.replace("<body>", "").replace("</body>", "")

        return html_content  # Send as raw HTML (not JSON)
    else:
        return "File not found", 404
    

@app.route('/generate_docx', methods=['POST'])
def generate_docx_file():
    """Generate a DOCX file from structured CV data."""
    data = request.json.get('structured_data', {})

    if not data:
        return jsonify({"error": "No structured data provided."}), 400

    # Generate DOCX
    output_docx_filename = "./uploads/output_cv.docx"
    generate_enhanced_docx(data, output_docx_filename)

    # Convert DOCX to HTML
    output_html_filename = "./uploads/output_cv.html"
    try:
        html_file = docx_to_html(output_docx_filename)
    except Exception as e:
        return jsonify({"error": f"Error converting DOCX to HTML: {str(e)}"}), 500

    # Return file links for both DOCX and HTML
    return jsonify({
        "message": "CV generated successfully", 
        "file": url_for('download_cv', filename='output_cv.docx'),
        "html_file": url_for('editor') 
    })

# @app.route('/download_html/<filename>')
# def download_html(filename):
#     return send_from_directory('uploads', filename)

@app.route('/uploads/<filename>')
def upload_file(filename):
    return send_from_directory('uploads', filename)

@app.route('/download_cv/<filename>')
def download_cv(filename):
    return send_from_directory('uploads', filename)

# -------------------- 🚀 RUN APP --------------------

if __name__ == '__main__':
    app.run(debug=True)
    