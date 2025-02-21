from flask import Flask, jsonify
import speech_recognition as sr
import pyttsx3
from docx import Document

app = Flask(__name__)

# Initialize speech recognizer and text-to-speech engine
recognizer = sr.Recognizer()
engine = pyttsx3.init()

# Function to ask questions using text-to-speech
def ask_question(question):
    engine.say(question)
    engine.runAndWait()

# Function to capture speech input and return the recognized text
def get_user_input(prompt):
    with sr.Microphone() as source:
        print(prompt)
        recognizer.adjust_for_ambient_noise(source)
        audio = recognizer.listen(source)

    user_input = ""
    user_confirmation = ""

    while True:
        try:
            # Recognize the speech using Google API
            user_input = recognizer.recognize_google(audio)
            print(f"You said: {user_input}, is that correct?")
            
            # Confirm the input with the user using TTS
            engine.say(f"You said: {user_input}, is that correct?")
            engine.runAndWait()

            # Listen for user confirmation
            with sr.Microphone() as source:
                audio = recognizer.listen(source)
                user_confirmation = recognizer.recognize_google(audio)
            
            # If user confirms, return the recognized input
            if user_confirmation.lower() in ["yes", "y"]:
                return user_input
            else:
                # If the user doesn't confirm, retry the question
                print("Please repeat your response.")
                engine.say("Please repeat your response.")
                engine.runAndWait()
                continue

        except sr.UnknownValueError:
            print("Sorry, I could not understand that. Can you say it again?")
            engine.say("Sorry, I could not understand that. Can you say it again?")
            engine.runAndWait()
            continue
        except sr.RequestError:
            print("Sorry, I'm having trouble with the speech recognition service.")
            engine.say("Sorry, I'm having trouble with the speech recognition service.")
            engine.runAndWait()
            return ""  # Retry on error

# Function to generate a DOCX file based on user input
def generate_docx(data, output_filename):
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)

    # Add user input sections to the DOCX file
    for section, content in data.items():
        doc.add_heading(section, level=1)
        for item in content:
            doc.add_paragraph(f"- {item}")

    # Save the document to the given filename
    doc.save(output_filename)
    print(f"Document saved as {output_filename}")

@app.route('/generate', methods=['GET'])
def generate_cv():
    # Define the questions and sections for the CV
    sections = {
        "Name": "What is your full name?",
        # "Email": "What is your email address?",
        # "Skills": "What are your key skills? Please list them.",
        # "Experience": "Please describe your work experience.",
        # "Education": "Where did you study and what degree did you earn?",
        # "Certification": "Do you have any certifications or awards?",
        # "References": "Please provide any references."
    }

    # Collect the data from the user via speech input
    cv_data = {}
    for section, question in sections.items():
        ask_question(question)  # Ask the user the question
        response = "Emmy"
        # get_user_input(f"Please answer the question: {question}")  # Get the response from the user
        cv_data[section] = response.split(",")  # Split by commas if multiple items (for example, skills)

    # Generate DOCX file with the collected data
    output_filename = "./output_cv.docx"
    generate_docx(cv_data, output_filename)

    return jsonify({"message": "CV generated successfully", "file": output_filename})

if __name__ == '__main__':
    app.run(debug=True)
